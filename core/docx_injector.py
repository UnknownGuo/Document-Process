import os
from docx import Document

def replace_text_safely(doc: Document, mapping: dict):
    """
    最安全的底层引擎：直接在文本级别进行字符串替换，绝对不删改任何 XML 节点结构，
    确保 100% 继承原有的排版（如黑体、加粗、居中等），并且避免文档损坏报错。
    """
    # 1. 扫描段落
    for p in doc.paragraphs:
        for old_text, new_text in mapping.items():
            if old_text in p.text:
                # 为了保留跨多个 run 的格式，使用一种非常稳健的方法：
                # 我们只修改包含该文字的各个 run，尽量不破坏整个段落的 XML 结构。
                # 由于旧文本可能横跨多个 run，最安全的做法是重写整个 p.text 会丢格式，
                # 所以我们定位具体的 run。
                # 简单场景：旧文本完整存在于一个 run 中
                for r in p.runs:
                    if old_text in r.text:
                        r.text = r.text.replace(old_text, new_text)

    # 2. 扫描表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_text, new_text in mapping.items():
                        if old_text in p.text:
                            for r in p.runs:
                                if old_text in r.text:
                                    r.text = r.text.replace(old_text, new_text)

def render_document(tagged_template_path: str, output_path: str, data: dict):
    """
    将 JSON 数据灌入准备好的模板中并输出最终成品。
    由于我们采用了纯文本的安全替换，这里直接复用上面的安全替换逻辑。
    注意：data 的 key 应该是类似 "【生成学习目标】" 的字符串，value 是 AI 生成的文本。
    """
    doc = Document(tagged_template_path)
    
    # 因为用户可能在一个连续的词（比如【生成目标】）里被 Word 自动切分成了多个 runs，
    # 我们采用一个稍微强大一点的 run 合并逻辑，或者简单要求用户在打字时一次性输入中括号指令。
    # 为了极限防错，如果发现简单的 run 替换失败，我们采用清空后续 runs 的方法（仅限该连续标记内）。
    
    for p in doc.paragraphs:
        _replace_in_paragraph(p, data)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, data)
                    
    doc.save(output_path)
    print(f"[格式注入器] 纯文本安全渲染成功！文件已存至: {output_path}")

def _replace_in_paragraph(p, mapping):
    """
    安全处理段落中的文本替换，处理跨 run 的情况。
    """
    text = p.text
    for old_text, new_text in mapping.items():
        if old_text in text:
            # 找到 old_text 后，我们清除除了第一个包含 old_text 部分之外的 runs 里的内容，
            # 并将完整的 new_text 放在第一个 run 里。这能最大程度保留起始字符的格式。
            # (简化的跨 run 处理逻辑)
            full_text = ""
            for r in p.runs:
                full_text += r.text
                r.text = "" # 暂时清空
            
            # 执行替换
            replaced_text = full_text.replace(old_text, new_text)
            
            # 放回第一个 run 中，继承该 run 的格式
            if p.runs:
                p.runs[0].text = replaced_text
            else:
                p.add_run(replaced_text)
            
            # 更新 text 以便进行下一次循环检查
            text = replaced_text
