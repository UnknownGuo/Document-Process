import os
import sys
import datetime
import json
from docx import Document
from dotenv import load_dotenv

# 环境初始化
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from core.llm_engine import generate_with_review
from qa_pipeline.continuous_learning import get_learning_context, save_lesson_learned
from qa_pipeline.proofreader import run_proofreader
from qa_pipeline.xml_diff import compare_xml_consistency
from qa_pipeline.visual_check import run_visual_layout_check

load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))

def safe_replace_block_text(block_element, new_text):
    """底层安全替换函数：清空文字并保留第一个 Run 的格式"""
    if hasattr(block_element, 'paragraphs'):
        if not block_element.paragraphs:
            block_element.add_paragraph(new_text)
            return
        first_p = block_element.paragraphs[0]
        if first_p.runs:
            first_run = first_p.runs[0]
            first_run.text = new_text
            for r in first_p.runs[1:]: r.text = ""
        else:
            first_p.add_run(new_text)
        for p in block_element.paragraphs[1:]:
            p_elem = p._p
            p_elem.getparent().remove(p_elem)
    else:
        if block_element.runs:
            first_run = block_element.runs[0]
            first_run.text = new_text
            for r in block_element.runs[1:]: r.text = ""
        else:
            block_element.add_run(new_text)

def run_test_pipeline(lesson_name: str):
    print(f"\n🚀 >>> 启动 AutoDoc-Agent V2.0 终极全流程管线: 《{lesson_name}》 <<<\n")
    
    template_path = "/mnt/a/AI_Studio/Auto-docs/老妈需求/template_with_tags.docx"
    output_dir = "/mnt/a/AI_Studio/Auto-docs/老妈需求/output"
    
    # 1. 强化学习记忆加载
    history_context = get_learning_context()
    
    # 2. 构造任务需求 (Context-Aware Parsing)
    doc = Document(template_path)
    blocks_to_process = {}
    
    # 段落解析
    for i, p in enumerate(doc.paragraphs):
        if "111" in p.text or "【" in p.text:
            blocks_to_process[f"p_{i}"] = {"context": "文档页眉标题区", "current_text": p.text}
            
    # 表格解析
    for t_idx, table in enumerate(doc.tables):
        processed_cells = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell._element in processed_cells: continue
                if "111" in cell.text or "【" in cell.text:
                    blocks_to_process[f"t_{t_idx}_r_{r_idx}_c_{c_idx}"] = {
                        "context": f"表格单元格(行{r_idx}, 列{c_idx})", "current_text": cell.text
                    }
                processed_cells.add(cell._element)

    # 3. Writer + Reviewer 协作生成
    # V2.1 修复：扁平化数据流
    # 我们把 context 信息写在 Prompt 文本里，但要求 AI 返回的 JSON 必须是纯净的 { "block_id": "纯中文" } 格式
    context_instructions = ""
    for b_id, b_data in blocks_to_process.items():
        context_instructions += f"- 【区块 ID: {b_id}】\n  - 位置上下文: {b_data['context']}\n  - 您的原始指令: {b_data['current_text']}\n\n"

    base_prompt = f"""
你是一位深耕部编版小学语文教材 20 年的特级教师。
请为六年级上册课文《{lesson_name}》生成教学设计中缺失的各个区块。
{history_context}

以下是你需要填写的区块列表及其上下文提示：
{context_instructions}

【绝对的红线要求：数据流扁平化】
你必须返回一个严格的 JSON。
这个 JSON 必须只有一层键值对！
Key 是上面列出的“区块 ID”（如 p_1, t_0_r_3_c_2）。
Value 必须是你想好的一段【纯中文字符串】（即直接生成的教案内容）。
绝对不允许在 Value 中嵌套字典！绝对不允许包含 "context" 或 "current_text" 等英文元数据！
正确示例：
{{
  "t_0_r_0_c_1": "《鲁滨逊漂流记》",
  "t_0_r_1_c_1": "1. 了解作者...\\n2. 学习手法..."
}}
"""
    
    constraints = """
1. 数量限制：学习目标和作业必须严格只有 2-3 条。
2. 指标格式：素养指标/任务动作必须分行显示，每行末尾紧跟一颗星 ⭐。
3. 身份固定：主备人王红晓，审核人为空。
4. 格式禁令：严禁使用自动编号（1. 2.），严禁生成文本表格，严禁添加书名号以外的前缀。
"""

    draft_data = generate_with_review(base_prompt, constraints)
    if not draft_data: return

    # 4. Proofreader 专业校对
    final_data = run_proofreader(draft_data)

    # 5. 无损格式渲染
    print("\n[Injector] 正在执行外科手术级无损排版注入...")
    # 重新加载 doc 准备注入
    doc = Document(template_path)
    processed_cells = set()
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                block_id = f"t_{t_idx}_r_{r_idx}_c_{c_idx}"
                if block_id in final_data and cell._element not in processed_cells:
                    # Fix: If the proofreader returned a dict instead of string, extract the text
                    text_to_insert = final_data[block_id]
                    if isinstance(text_to_insert, dict):
                        # Attempt to extract just the generated text (it might be under current_text or just use values)
                        text_to_insert = text_to_insert.get("current_text", str(text_to_insert))
                    elif isinstance(text_to_insert, list):
                        text_to_insert = "\n".join(str(x) for x in text_to_insert)
                    else:
                        text_to_insert = str(text_to_insert)
                    
                    safe_replace_block_text(cell, text_to_insert)
                    processed_cells.add(cell._element)
    for i, p in enumerate(doc.paragraphs):
        block_id = f"p_{i}"
        if block_id in final_data:
            text_to_insert = final_data[block_id]
            if isinstance(text_to_insert, dict):
                text_to_insert = text_to_insert.get("current_text", str(text_to_insert))
            elif isinstance(text_to_insert, list):
                text_to_insert = "\n".join(str(x) for x in text_to_insert)
            else:
                text_to_insert = str(text_to_insert)
            safe_replace_block_text(p, text_to_insert)

    today = datetime.datetime.now().strftime("%Y%m%d")
    final_docx_path = os.path.join(output_dir, f"{lesson_name}_v8_FinalQA_{today}.docx")
    doc.save(final_docx_path)
    print(f"  -> 成品已保存: {final_docx_path}")

    # 6. XML Deep Diff 格式完整性校验
    success, msg = compare_xml_consistency(template_path, final_docx_path)
    
    # 7. (可选) 视觉校验
    # template_screenshot = "/mnt/a/AI_Studio/Auto-docs/老妈需求/template_snapshot.png"
    # output_screenshot = "/mnt/a/AI_Studio/Auto-docs/老妈需求/output/generated_snapshot.png"
    # run_visual_layout_check(template_screenshot, output_screenshot)

    print("\n✅ 全流程闭环测试完成。")

if __name__ == "__main__":
    run_test_pipeline("鲁滨逊漂流记")
