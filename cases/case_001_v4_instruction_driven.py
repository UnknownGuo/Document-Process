import os
import sys
import datetime
import re
from docx import Document
from dotenv import load_dotenv

# 确保能导入 core 模块
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.llm_engine import generate_structured_data
from core.docx_injector import render_document

load_dotenv(os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env'))

def run_instruction_case(lesson_name: str):
    print(f"========== 开始执行指令驱动案例: {lesson_name} ==========")
    template_path = "/mnt/a/AI_Studio/Auto-docs/老妈需求/template_with_tags.docx"
    output_dir = "/mnt/a/AI_Studio/Auto-docs/老妈需求/output"
    
    if not os.path.exists(template_path):
        print(f"错误：找不到模板文件 {template_path}")
        return

    doc = Document(template_path)
    
    # 1. 提取所有唯一指令
    instruction_markers = []
    
    def extract_from_paragraphs(paragraphs):
        found = []
        for p in paragraphs:
            found.extend(re.findall(r"【(.*?)】", p.text))
        return found

    instruction_markers.extend(extract_from_paragraphs(doc.paragraphs))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                instruction_markers.extend(extract_from_paragraphs(cell.paragraphs))
    
    unique_instructions = list(dict.fromkeys(instruction_markers)) # 保持顺序去重
    print(f"检测到 {len(unique_instructions)} 个不同的中文指令。")

    # 2. 构造 Prompt 让 AI 填充这些指令
    # 我们将所有指令编号，让 AI 返回一个对应的 JSON
    instructions_formatted = "\n".join([f"{i+1}. 【{inst}】" for i, inst in enumerate(unique_instructions)])
    
    prompt = f"""
你是一位深耕部编版小学语文教材 20 年的特级教师。
目前正在为六年级上册第二单元课文《{lesson_name}》编写教学设计。

请根据以下模板中提取出的【中文指令】，提供对应的填空内容。
要求：
1. 内容精炼、专业，符合小学语文教学规范。
2. 严禁输出多余的回车或符号。
3. 如果指令中提到“加星星”，请在内容末尾加上“达成度：⭐⭐⭐⭐”。
4. 必须严格返回 JSON 格式，Key 是指令的序号（字符串），Value 是生成的内容。

待处理指令列表：
{instructions_formatted}
"""

    print("正在请求 Gemini 2.5 Pro 生成内容...")
    result_data = generate_structured_data(prompt, model_name='gemini-2.5-pro')
    
    if not result_data:
        print("AI 生成失败。")
        return

    # 3. 构造替换映射表：把“【指令】”映射到“生成的内容”
    mapping = {}
    for i, inst in enumerate(unique_instructions):
        key = str(i + 1)
        if key in result_data:
            mapping[f"【{inst}】"] = result_data[key]
        else:
            mapping[f"【{inst}】"] = "（未生成内容）"

    # 特殊处理：年级标题（如果没在指令里，我们手动补一个，或者让用户写在指令里）
    # 在这个模板里，用户可能直接写了“年级语文...”，我们需要保持
    
    # 4. 执行安全渲染
    today = datetime.datetime.now().strftime("%Y%m%d")
    out_path = os.path.join(output_dir, f"{lesson_name}_v4_{today}.docx")
    
    render_document(template_path, out_path, mapping)
    print(f"已生成并保存：{out_path}")

if __name__ == "__main__":
    # 先跑一个测试
    run_instruction_case("鲁滨逊漂流记")
