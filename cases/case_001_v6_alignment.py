import os
import sys
import datetime
import json
from docx import Document
from dotenv import load_dotenv

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.llm_engine import generate_structured_data

load_dotenv(os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env'))

def get_table_headers(table, r_idx, c_idx):
    try:
        col_header = table.rows[0].cells[c_idx].text.strip().replace("\n", " ")
    except:
        col_header = ""
    try:
        row_header = table.rows[r_idx].cells[0].text.strip().replace("\n", " ")
    except:
        row_header = ""
    return col_header, row_header

def parse_document_blocks(doc_path):
    doc = Document(doc_path)
    blocks_to_process = {}
    
    for i, p in enumerate(doc.paragraphs):
        if "111" in p.text or "【" in p.text:
            blocks_to_process[f"p_{i}"] = {
                "type": "paragraph",
                "context": "文档顶部区域",
                "current_text": p.text
            }
            
    for t_idx, table in enumerate(doc.tables):
        processed_cells = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = cell._element
                if cell_id in processed_cells:
                    continue
                
                text = cell.text
                if "111" in text or "【" in text:
                    col_h, row_h = get_table_headers(table, r_idx, c_idx)
                    blocks_to_process[f"t_{t_idx}_r_{r_idx}_c_{c_idx}"] = {
                        "type": "table_cell",
                        "context": f"行表头：{row_h} | 列表头：{col_h}",
                        "current_text": text
                    }
                processed_cells.add(cell_id)
                
    return doc, blocks_to_process

def safe_replace_block_text(block_element, new_text):
    if hasattr(block_element, 'paragraphs'):
        if not block_element.paragraphs:
            block_element.add_paragraph(new_text)
            return
        first_p = block_element.paragraphs[0]
        if first_p.runs:
            first_run = first_p.runs[0]
            first_run.text = new_text
            for r in first_p.runs[1:]: r.text = ""
        else:
            first_p.add_run(new_text)
        for p in block_element.paragraphs[1:]:
            p_elem = p._p
            p_elem.getparent().remove(p_elem)
    else:
        if block_element.runs:
            first_run = block_element.runs[0]
            first_run.text = new_text
            for r in block_element.runs[1:]: r.text = ""
        else:
            block_element.add_run(new_text)

def run_alignment_case(lesson_name: str, author: str = "王红晓"):
    print(f"========== 开始执行 V6 意图对齐版生成: {lesson_name} ==========")
    template_path = "/mnt/a/AI_Studio/Auto-docs/老妈需求/template_with_tags.docx"
    output_dir = "/mnt/a/AI_Studio/Auto-docs/老妈需求/output"
    
    doc, blocks = parse_document_blocks(template_path)
    if not blocks:
        print("未在模板中检测到任何包含 111 或 【】 的标记。")
        return

    blocks_json_str = json.dumps(blocks, ensure_ascii=False, indent=2)
    
    prompt = f"""
你是一位深耕部编版小学语文教材 20 年的特级教师。
当前任务：为六年级上册课文《{lesson_name}》编写教学设计。

用户提供了一个含有旧课文（《记金华的双龙洞》）示例内容的模板。
用户在需要你根据新课文（《{lesson_name}》）重新生成的地方，打上了诸如“1111”、“【中文指令】”的标记。

下面是一个 JSON，列出了所有检测到标记的区块 (Block)。
你的任务是：为每一个 Block 生成【全新的文本】，完全替换掉原有的示例文本和标记。

核心【意图对齐】要求（请务必严格遵守）：
1. 【极简课题】：课题处只输出书名（如《{lesson_name}》），严禁带有类似“快乐读书吧：”等前缀。
2. 【数量精准控制】：“学习目标”和“作业”栏目，必须严格只写 2-3 条。不多不少。
3. 【原子化评价指标】：对于“素养指标”或任务左侧的概括词（如原文的“说游踪 理顺序”），必须将每个动作拆分成独立的一行，并在每一行的末尾单独跟上一颗星星 ⭐。
   正确格式示例：
   知背景 ⭐
   理情节 ⭐
   谈初感 ⭐
   （绝对不允许挤在同一行或在最后统一写“达成度：⭐⭐⭐”）
4. 【主备人】：主备人固定为“{author}”，审核人为空。
5. 【纯文本格式】：坚决不使用 Word 自动编号（如1. 2.），请直接输出纯文本段落。坚决不要用文本画表格。

返回格式：
请返回一个严格的 JSON。Key 是传入的 Block ID（如 "t_0_r_3_c_2"），Value 是你为该区块生成的全新内容。

待处理区块数据：
{blocks_json_str}
"""
    
    print("正在请求 Gemini 2.5 Pro 融合您的修改意图进行重写...")
    result_data = generate_structured_data(prompt, model_name='gemini-2.5-pro')
    
    if not result_data:
        print("AI 生成失败。")
        return
        
    print("开始进行意图对齐替换...")
    processed_cells = set()
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id_str = f"t_{t_idx}_r_{r_idx}_c_{c_idx}"
                if cell_id_str in result_data:
                    cell_elem_id = cell._element
                    if cell_elem_id not in processed_cells:
                        new_text = result_data[cell_id_str]
                        safe_replace_block_text(cell, new_text)
                        processed_cells.add(cell_elem_id)
                        
    for i, p in enumerate(doc.paragraphs):
        p_id_str = f"p_{i}"
        if p_id_str in result_data:
            new_text = result_data[p_id_str]
            safe_replace_block_text(p, new_text)
            
    today = datetime.datetime.now().strftime("%Y%m%d")
    out_path = os.path.join(output_dir, f"{lesson_name}_v6_{today}.docx")
    
    doc.save(out_path)
    print(f"✅ 已生成对齐您修改意图的终极版：{out_path}")

if __name__ == "__main__":
    run_alignment_case("鲁滨逊漂流记")
