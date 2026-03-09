import os
import sys
import datetime
import json
from docx import Document
from dotenv import load_dotenv

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.llm_engine import generate_structured_data

load_dotenv(os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env'))

def get_table_headers(table, r_idx, c_idx):
    try:
        col_header = table.rows[0].cells[c_idx].text.strip().replace("\n", " ")
    except:
        col_header = ""
    try:
        row_header = table.rows[r_idx].cells[0].text.strip().replace("\n", " ")
    except:
        row_header = ""
    return col_header, row_header

def parse_document_blocks(doc_path):
    doc = Document(doc_path)
    blocks_to_process = {}
    
    # 1. Paragraphs
    for i, p in enumerate(doc.paragraphs):
        if "111" in p.text or "【" in p.text:
            blocks_to_process[f"p_{i}"] = {
                "type": "paragraph",
                "context": "文档顶部区域",
                "current_text": p.text
            }
            
    # 2. Tables
    for t_idx, table in enumerate(doc.tables):
        # Keep track of processed cells to handle merged cells
        processed_cells = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = cell._element
                if cell_id in processed_cells:
                    continue
                
                text = cell.text
                if "111" in text or "【" in text:
                    col_h, row_h = get_table_headers(table, r_idx, c_idx)
                    blocks_to_process[f"t_{t_idx}_r_{r_idx}_c_{c_idx}"] = {
                        "type": "table_cell",
                        "context": f"行表头：{row_h} | 列表头：{col_h}",
                        "current_text": text
                    }
                processed_cells.add(cell_id)
                
    return doc, blocks_to_process

def safe_replace_block_text(block_element, new_text):
    """Safely clear block element (Paragraph or Cell) and insert new text retaining formatting of the first run."""
    if hasattr(block_element, 'paragraphs'):
        # It's a cell
        if not block_element.paragraphs:
            block_element.add_paragraph(new_text)
            return
        
        first_p = block_element.paragraphs[0]
        if first_p.runs:
            first_run = first_p.runs[0]
            first_run.text = new_text
            for r in first_p.runs[1:]: r.text = ""
        else:
            first_p.add_run(new_text)
            
        for p in block_element.paragraphs[1:]:
            p_elem = p._p
            p_elem.getparent().remove(p_elem)
            
    else:
        # It's a paragraph
        if block_element.runs:
            first_run = block_element.runs[0]
            first_run.text = new_text
            for r in block_element.runs[1:]: r.text = ""
        else:
            block_element.add_run(new_text)

def run_context_aware_case(lesson_name: str):
    print(f"========== 开始执行 Context-Aware 生成案例: {lesson_name} ==========")
    template_path = "/mnt/a/AI_Studio/Auto-docs/老妈需求/template_with_tags.docx"
    output_dir = "/mnt/a/AI_Studio/Auto-docs/老妈需求/output"
    
    doc, blocks = parse_document_blocks(template_path)
    
    if not blocks:
        print("未在模板中检测到任何包含 111 或 【】 的标记。")
        return

    print(f"检测到 {len(blocks)} 个待生成区块。")
    
    blocks_json_str = json.dumps(blocks, ensure_ascii=False, indent=2)
    
    prompt = f"""
你是一位深耕部编版小学语文教材 20 年的特级教师。
当前任务：为六年级上册课文《{lesson_name}》编写教学设计。

用户提供了一个含有旧课文（《记金华的双龙洞》）示例内容的模板。
用户在需要你根据新课文（《{lesson_name}》）重新生成的地方，打上了诸如“1111”、“111”、“【中文指令】”的标记。

下面是一个 JSON，列出了所有检测到标记的区块 (Block)。
你的任务是：为每一个 Block 生成【全新的文本】，完全替换掉原有的示例文本和标记。

核心要求：
1. 【极简、精准】：参考原示例文本的风格，但字数减半。绝不要长篇大论！坚决不使用 Word 自动编号（如1. 2.），用纯文本。坚决不输出多余的空行。
2. 【星星评价】：如果原文中提到了星标或指令要求加星星，请在生成的文本末尾直接加上纯文本星星，例如：“达成度：⭐⭐⭐⭐”。
3. 【理解上下文】：请根据每个 Block 的 `context`（行/列表头）判断这里该填什么（如：学习目标、活动过程、板书、作业等）。
4. 【清理大括号】：如果用户的指令中写了“和{{}}类似，但是你要删去{{}}”，这意味着你要提取该部分原有的结构逻辑，但不需要输出大括号。

返回格式：
请返回一个严格的 JSON。
Key 是传入的 Block ID（如 "t_0_r_3_c_2"），Value 是你为该区块生成的【全新纯文本内容】。

待处理区块数据：
{blocks_json_str}
"""
    
    print("正在请求 Gemini 2.5 Pro 根据上下文进行智能重写...")
    result_data = generate_structured_data(prompt, model_name='gemini-2.5-pro')
    
    if not result_data:
        print("AI 生成失败。")
        return
        
    print("开始进行外科手术级精准替换...")
    # Map generated text back to the document
    processed_cells = set()
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id_str = f"t_{t_idx}_r_{r_idx}_c_{c_idx}"
                if cell_id_str in result_data:
                    cell_elem_id = cell._element
                    if cell_elem_id not in processed_cells:
                        new_text = result_data[cell_id_str]
                        safe_replace_block_text(cell, new_text)
                        processed_cells.add(cell_elem_id)
                        
    for i, p in enumerate(doc.paragraphs):
        p_id_str = f"p_{i}"
        if p_id_str in result_data:
            new_text = result_data[p_id_str]
            safe_replace_block_text(p, new_text)
            
    today = datetime.datetime.now().strftime("%Y%m%d")
    out_path = os.path.join(output_dir, f"{lesson_name}_v5_{today}.docx")
    
    doc.save(out_path)
    print(f"已生成完美格式的教案：{out_path}")

if __name__ == "__main__":
    run_context_aware_case("鲁滨逊漂流记")
