import os
import sys
import datetime
import json
from docx import Document
from dotenv import load_dotenv

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.llm_engine import generate_with_review

load_dotenv(os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env'))

def get_table_headers(table, r_idx, c_idx):
    try:
        col_header = table.rows[0].cells[c_idx].text.strip().replace("\n", " ")
    except:
        col_header = ""
    try:
        row_header = table.rows[r_idx].cells[0].text.strip().replace("\n", " ")
    except:
        row_header = ""
    return col_header, row_header

def parse_document_blocks(doc_path):
    doc = Document(doc_path)
    blocks_to_process = {}
    
    for i, p in enumerate(doc.paragraphs):
        if "111" in p.text or "【" in p.text:
            blocks_to_process[f"p_{i}"] = {"type": "paragraph", "context": "文档顶部区域", "current_text": p.text}
            
    for t_idx, table in enumerate(doc.tables):
        processed_cells = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = cell._element
                if cell_id in processed_cells: continue
                text = cell.text
                if "111" in text or "【" in text:
                    col_h, row_h = get_table_headers(table, r_idx, c_idx)
                    blocks_to_process[f"t_{t_idx}_r_{r_idx}_c_{c_idx}"] = {
                        "type": "table_cell", "context": f"行表头：{row_h} | 列表头：{col_h}", "current_text": text
                    }
                processed_cells.add(cell_id)
                
    return doc, blocks_to_process

def safe_replace_block_text(block_element, new_text):
    if hasattr(block_element, 'paragraphs'):
        if not block_element.paragraphs:
            block_element.add_paragraph(new_text)
            return
        first_p = block_element.paragraphs[0]
        if first_p.runs:
            first_run = first_p.runs[0]
            first_run.text = new_text
            for r in first_p.runs[1:]: r.text = ""
        else:
            first_p.add_run(new_text)
        for p in block_element.paragraphs[1:]:
            p_elem = p._p
            p_elem.getparent().remove(p_elem)
    else:
        if block_element.runs:
            first_run = block_element.runs[0]
            first_run.text = new_text
            for r in block_element.runs[1:]: r.text = ""
        else:
            block_element.add_run(new_text)

def run_hybrid_batch_case():
    print("========== 开始执行 V7 混合架构(Writer+Reviewer) 批量生成 ==========")
    template_path = "/mnt/a/AI_Studio/Auto-docs/老妈需求/template_with_tags.docx"
    output_dir = "/mnt/a/AI_Studio/Auto-docs/老妈需求/output"
    os.makedirs(output_dir, exist_ok=True)
    
    lessons = [
        "骑鹅历险记",
        "汤姆索亚历险记",
        "口语交际：通读一本书",
        "习作：写作梗概",
        "语文园地",
        "快乐读书吧"
    ]
    
    author = "王红晓"
    
    for lesson_name in lessons:
        print(f"\n>>>> 正在处理课文：《{lesson_name}》 <<<<")
        doc, blocks = parse_document_blocks(template_path)
        if not blocks:
            print(f"模板解析失败。跳过《{lesson_name}》。")
            continue

        blocks_json_str = json.dumps(blocks, ensure_ascii=False, indent=2)
        
        # 1. 定义写作者的 Prompt
        base_prompt = f"""
你是一位深耕部编版小学语文教材 20 年的特级教师。
当前任务：为六年级上册课文《{lesson_name}》编写教学设计。

用户提供了一个含有旧内容的模板，且在需要你重新生成的地方打上了标记（如 1111 或 【】）。
请根据下面 JSON 列出的待处理区块，为每一个 Block 生成【全新的文本】，完全替换掉原有的示例文本。

请返回一个严格的 JSON。Key 是 Block ID，Value 是生成的全新内容。

待处理区块数据：
{blocks_json_str}
"""
        
        # 2. 定义严格的质检标准 (Constraints)
        constraints = f"""
1. 【极简课题】：课题处只输出书名（如《{lesson_name}》），严禁带有类似“快乐读书吧：”等前缀。
2. 【数量精准控制】：“学习目标”和“作业”栏目，必须严格只写 2-3 条。不能只有 1 条，也不能多于 3 条。
3. 【原子化评价指标】：对于“素养指标”或任务左侧的概括词（如：说游踪、理顺序），必须拆分成独立的行，并在每一行的末尾单独跟上一颗星星 ⭐。
   【必须符合这种格式】：
   知背景 ⭐
   理情节 ⭐
   谈初感 ⭐
   （绝对不允许挤在同一行，绝对不允许在结尾统一写“达成度：⭐⭐⭐”）
4. 【主备人】：主备人固定为“{author}”，审核人为空。
5. 【纯文本格式】：坚决不使用 Word 自动编号（如1. 2.），直接输出纯文本段落。坚决不要用虚线画表格。
"""

        # 3. 调用多智能体管线：生成 -> 审查 -> (可能)返工 -> 最终输出
        result_data = generate_with_review(base_prompt, constraints, model_name='gemini-2.5-pro')
        
        if not result_data:
            print(f"❌ AI 生成失败，跳过《{lesson_name}》。")
            continue
            
        print("[格式注入器] 审查通过！开始进行外科手术级精准替换...")
        processed_cells = set()
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_id_str = f"t_{t_idx}_r_{r_idx}_c_{c_idx}"
                    if cell_id_str in result_data:
                        cell_elem_id = cell._element
                        if cell_elem_id not in processed_cells:
                            new_text = result_data[cell_id_str]
                            safe_replace_block_text(cell, new_text)
                            processed_cells.add(cell_elem_id)
                            
        for i, p in enumerate(doc.paragraphs):
            p_id_str = f"p_{i}"
            if p_id_str in result_data:
                new_text = result_data[p_id_str]
                safe_replace_block_text(p, new_text)
                
        today = datetime.datetime.now().strftime("%Y%m%d")
        out_path = os.path.join(output_dir, f"{lesson_name}_v7_AutoReview_{today}.docx")
        
        doc.save(out_path)
        print(f"✅ 《{lesson_name}》生成完毕，已保存至：{out_path}")

if __name__ == "__main__":
    run_hybrid_batch_case()
